"""Extract plain text from Word (.docx) briefs for the harness planner."""
from __future__ import annotations

from pathlib import Path

MAX_DOC_CHARS = 48_000
ALLOWED_EXTENSIONS = {".docx"}


def allowed_document_name(filename: str) -> bool:
    if not filename:
        return False
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from e

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    body = "\n\n".join(parts).strip()
    if not body:
        raise ValueError("The Word document appears empty — add headings and paragraphs first.")

    if len(body) > MAX_DOC_CHARS:
        body = (
            body[:MAX_DOC_CHARS]
            + "\n\n[Document truncated for planner context — full file kept in workspace/uploads/]"
        )
    return body


def merge_brief(user_idea: str, document_text: str, document_name: str = "") -> str:
    """Combine optional typed notes with extracted document text."""
    sections: list[str] = []
    doc = (document_text or "").strip()
    notes = (user_idea or "").strip()
    label = document_name or "uploaded document"

    if doc:
        sections.append(f"## Source document ({label})\n\n{doc}")
    if notes:
        sections.append(f"## Additional instructions\n\n{notes}")
    return "\n\n".join(sections)


def persist_brief_markdown(
    run_dir: Path, document_text: str, document_name: str, user_idea: str = ""
) -> None:
    mem = run_dir / "memory"
    mem.mkdir(parents=True, exist_ok=True)
    combined = merge_brief(user_idea, document_text, document_name)
    (mem / "SOURCE_BRIEF.md").write_text(
        f"# Project brief\n\n{combined}\n", encoding="utf-8"
    )
