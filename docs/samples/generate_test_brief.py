"""Generate manager test brief (.docx) for IHS upload. Run from repo root:
  cd backend && .venv\\Scripts\\activate && python ..\\docs\\samples\\generate_test_brief.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "IHS-Test-Brief-Simple-Web-App.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def main() -> None:
    doc = Document()

    title = doc.add_heading("Project Brief — Simple Web App (Test)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run("Document type: ").bold = True
    meta.add_run("Manager acceptance test for Intelligent Harness System\n")
    meta.add_run("Target stack: ").bold = True
    meta.add_run("HTML, CSS, and JavaScript only (no React, no build tools)\n")
    meta.add_run("Suggested title: ").bold = True
    meta.add_run("Focus Timer — Pomodoro Style")

    add_heading(doc, "1. Summary", 1)
    doc.add_paragraph(
        "Build a small, single-page web application that helps a user run focused "
        "work sessions using a Pomodoro-style timer. The app must run in the browser "
        "using only static HTML, CSS, and vanilla JavaScript. No frameworks, no npm, "
        "and no backend server."
    )

    add_heading(doc, "2. Goals", 1)
    add_bullet(doc, "Let the user start, pause, and reset a countdown timer.")
    add_bullet(doc, "Offer preset durations: 25 minutes (focus) and 5 minutes (break).")
    add_bullet(doc, "Show remaining time clearly (minutes and seconds).")
    add_bullet(doc, "Play a short visual cue when the timer reaches zero (color flash or message).")
    add_bullet(doc, "Work on desktop and mobile (responsive layout).")

    add_heading(doc, "3. Pages & layout", 1)
    doc.add_paragraph("One page (index.html) with these sections:")
    add_bullet(doc, "Header: app name “Focus Timer” and one-line tagline.")
    add_bullet(doc, "Main: large digital clock display (MM:SS).")
    add_bullet(doc, "Controls: Start, Pause, Reset buttons.")
    add_bullet(doc, "Presets: two buttons — “25 min Focus” and “5 min Break”.")
    add_bullet(doc, "Footer: text “Built with HTML, CSS & JavaScript”.")

    add_heading(doc, "4. Functional requirements", 1)
    add_bullet(doc, "Default duration on load: 25:00 (focus mode).")
    add_bullet(doc, "Start begins counting down every second; Pause stops the interval.")
    add_bullet(doc, "Reset returns to the currently selected preset duration.")
    add_bullet(doc, "Clicking a preset updates the display and resets the timer to that duration.")
    add_bullet(doc, "At 00:00, show “Time’s up!” and highlight the timer area (CSS class).")
    add_bullet(doc, "Prevent starting a second interval if one is already running.")

    add_heading(doc, "5. Design requirements", 1)
    add_bullet(doc, "Clean, modern look: centered card on a soft gradient background.")
    add_bullet(doc, "Readable font (system font stack is fine).")
    add_bullet(doc, "Primary button color: teal or blue; hover states on buttons.")
    add_bullet(doc, "Timer digits at least 48px on desktop.")
    add_bullet(doc, "Use CSS variables for colors so the theme is easy to tweak.")

    add_heading(doc, "6. Technical constraints", 1)
    add_bullet(doc, "Files: index.html, styles.css, script.js (linked from HTML).")
    add_bullet(doc, "No external libraries (no jQuery, no React, no CDN frameworks).")
    add_bullet(doc, "JavaScript must use setInterval or requestAnimationFrame responsibly; clear interval on pause/reset.")
    add_bullet(doc, "Semantic HTML5 (header, main, footer, button elements).")
    add_bullet(doc, "All styling in styles.css; no inline styles except none required.")

    add_heading(doc, "7. Acceptance criteria (for manager sign-off)", 1)
    criteria = [
        "Opening index.html in Chrome shows the timer at 25:00.",
        "Start counts down; Pause freezes the display; Reset restores the preset.",
        "“5 min Break” switches display to 05:00 and reset works.",
        "At zero, “Time’s up!” appears and the timer area is visually highlighted.",
        "Layout looks acceptable on a phone-width viewport (no horizontal scroll).",
        "No console errors during normal use.",
    ]
    for i, item in enumerate(criteria, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")

    add_heading(doc, "8. Out of scope", 1)
    add_bullet(doc, "User accounts, databases, or API calls.")
    add_bullet(doc, "Sound effects (optional nice-to-have only if trivial).")
    add_bullet(doc, "PWA, service workers, or install prompts.")

    add_heading(doc, "9. How to test in IHS", 1)
    steps = [
        "Open the IHS dashboard (live demo URL from your team).",
        "Go to New Project.",
        "Upload this Word file as the project brief (.docx).",
        "Optional: add a short note in the text box: “Static HTML/CSS/JS Pomodoro timer”.",
        "Start the harness run and follow Logs → Debate → Approval → Build.",
        "After build, open the generated app and verify acceptance criteria above.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    doc.add_paragraph()
    p = doc.add_paragraph("End of brief.")
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
